# coding=utf-8
"""
    @project: maxkb
    @Author：虎
    @file： text_split_handle.py
    @date：2024/3/27 18:19
    @desc:
"""
import io
import os
import re
import traceback
import uuid
from functools import reduce
from typing import List

from docx import Document, ImagePart
from docx.oxml import ns
from docx.table import Table
from docx.text.paragraph import Paragraph

from common.handle.base_split_handle import BaseSplitHandle
from common.util.split_model import SplitModel
from dataset.models import Image
from django.utils.translation import gettext_lazy as _

default_pattern_list = [re.compile('(?<=^)# .*|(?<=\\n)# .*'),
                        re.compile('(?<=\\n)(?<!#)## (?!#).*|(?<=^)(?<!#)## (?!#).*'),
                        re.compile("(?<=\\n)(?<!#)### (?!#).*|(?<=^)(?<!#)### (?!#).*"),
                        re.compile("(?<=\\n)(?<!#)#### (?!#).*|(?<=^)(?<!#)#### (?!#).*"),
                        re.compile("(?<=\\n)(?<!#)##### (?!#).*|(?<=^)(?<!#)##### (?!#).*"),
                        re.compile("(?<=\\n)(?<!#)###### (?!#).*|(?<=^)(?<!#)###### (?!#).*")]

old_docx_nsmap = {'v': 'urn:schemas-microsoft-com:vml'}
combine_nsmap = {**ns.nsmap, **old_docx_nsmap}


def image_to_mode(image, doc: Document, images_list, get_image_id):
    image_ids = image['get_image_id_handle'](image.get('image'))
    for img_id in image_ids:  # 获取图片id
        part = doc.part.related_parts[img_id]  # 根据图片id获取对应的图片
        if isinstance(part, ImagePart):
            image_uuid = get_image_id(img_id)
            if len([i for i in images_list if i.id == image_uuid]) == 0:
                image = Image(id=image_uuid, image=part.blob, image_name=part.filename)
                images_list.append(image)
            return f'![](/api/image/{image_uuid})'


def get_paragraph_element_images(paragraph_element, doc: Document, images_list, get_image_id):
    images_xpath_list = [(".//pic:pic", lambda img: img.xpath('.//a:blip/@r:embed')),
                         (".//w:pict", lambda img: img.xpath('.//v:imagedata/@r:id', namespaces=combine_nsmap))]
    images = []
    for images_xpath, get_image_id_handle in images_xpath_list:
        try:
            _images = paragraph_element.xpath(images_xpath)
            if _images is not None and len(_images) > 0:
                for image in _images:
                    images.append({'image': image, 'get_image_id_handle': get_image_id_handle})
        except Exception as e:
            pass
    return images


def images_to_string(images, doc: Document, images_list, get_image_id):
    return "".join(
        [item for item in [image_to_mode(image, doc, images_list, get_image_id) for image in images] if
         item is not None])


def get_paragraph_element_txt(paragraph_element, doc: Document, images_list, get_image_id):
    try:
        images = get_paragraph_element_images(paragraph_element, doc, images_list, get_image_id)
        if len(images) > 0:
            return images_to_string(images, doc, images_list, get_image_id)
        elif paragraph_element.text is not None:
            return paragraph_element.text
        return ""
    except Exception as e:
        print(e)
    return ""


def get_paragraph_txt(paragraph: Paragraph, doc: Document, images_list, get_image_id):
    try:
        return "".join([get_paragraph_element_txt(e, doc, images_list, get_image_id) for e in paragraph._element])
    except Exception as e:
        return ""


def get_cell_text(cell, doc: Document, images_list, get_image_id):
    try:
        return "".join(
            [get_paragraph_txt(paragraph, doc, images_list, get_image_id) for paragraph in cell.paragraphs]).replace(
            "\n", '</br>')
    except Exception as e:
        return ""


def get_image_id_func():
    image_map = {}

    def get_image_id(image_id):
        _v = image_map.get(image_id)
        if _v is None:
            image_map[image_id] = uuid.uuid1()
            return image_map.get(image_id)
        return _v

    return get_image_id


title_font_list = [
    [36, 100],
    [30, 36]
]


def get_title_level(paragraph: Paragraph):
    try:
        if paragraph.style is not None:
            psn = paragraph.style.name
            if psn.startswith('Heading') or psn.startswith('TOC 标题') or psn.startswith('标题'):
                return int(psn.replace("Heading ", '').replace('TOC 标题', '').replace('标题',
                                                                                       ''))
        if len(paragraph.runs) == 1:
            font_size = paragraph.runs[0].font.size
            pt = font_size.pt
            if pt >= 30:
                for _value, index in zip(title_font_list, range(len(title_font_list))):
                    if pt >= _value[0] and pt < _value[1]:
                        return index + 1
    except Exception as e:
        pass
    return None


class DocSplitHandle(BaseSplitHandle):
    @staticmethod
    def paragraph_to_md(paragraph: Paragraph, doc: Document, images_list, get_image_id):
        try:
            title_level = get_title_level(paragraph)
            if title_level is not None:
                title = "".join(["#" for i in range(title_level)]) + " " + paragraph.text
                images = reduce(lambda x, y: [*x, *y],
                                [get_paragraph_element_images(e, doc, images_list, get_image_id) for e in
                                 paragraph._element],
                                [])
                if len(images) > 0:
                    return title + '\n' + images_to_string(images, doc, images_list, get_image_id) if len(
                        paragraph.text) > 0 else images_to_string(images, doc, images_list, get_image_id)
                return title

        except Exception as e:
            traceback.print_exc()
            return paragraph.text
        return get_paragraph_txt(paragraph, doc, images_list, get_image_id)

    @staticmethod
    def table_to_md(table, doc: Document, images_list, get_image_id):
        rows = table.rows

        # 创建 Markdown 格式的表格
        md_table = '| ' + ' | '.join(
            [get_cell_text(cell, doc, images_list, get_image_id) for cell in rows[0].cells]) + ' |\n'
        md_table += '| ' + ' | '.join(['---' for i in range(len(rows[0].cells))]) + ' |\n'
        for row in rows[1:]:
            md_table += '| ' + ' | '.join(
                [get_cell_text(cell, doc, images_list, get_image_id) for cell in row.cells]) + ' |\n'
        return md_table

    def to_md(self, doc, images_list, get_image_id):
        elements = []
        for element in doc.element.body:
            tag = str(element.tag)
            if tag.endswith('tbl'):
                # 处理表格
                table = Table(element, doc)
                elements.append(table)
            elif tag.endswith('p'):
                # 处理段落
                paragraph = Paragraph(element, doc)
                elements.append(paragraph)
        return "\n".join(
            [self.paragraph_to_md(element, doc, images_list, get_image_id) if isinstance(element,
                                                                                         Paragraph) else self.table_to_md(
                element,
                doc,
                images_list, get_image_id)
             for element
             in elements])

    def handle(self, file, pattern_list: List, with_filter: bool, limit: int, get_buffer, save_image):
        file_name = os.path.basename(file.name)
        try:
            image_list = []
            buffer = get_buffer(file)
            doc = Document(io.BytesIO(buffer))
            content = self.to_md(doc, image_list, get_image_id_func())
            if len(image_list) > 0:
                save_image(image_list)
            if pattern_list is not None and len(pattern_list) > 0:
                split_model = SplitModel(pattern_list, with_filter, limit)
            else:
                split_model = SplitModel(default_pattern_list, with_filter=with_filter, limit=limit)
        except BaseException as e:
            traceback.print_exception(e)
            return {'name': file_name,
                    'content': []}
        return {'name': file_name,
                'content': split_model.parse(content)
                }

    def support(self, file, get_buffer):
        file_name: str = file.name.lower()
        if file_name.endswith(".docx") or file_name.endswith(".doc") or file_name.endswith(
                ".DOC") or file_name.endswith(".DOCX"):
            return True
        return False

    def get_content(self, file, save_image):
        try:
            image_list = []
            buffer = file.read()
            doc = Document(io.BytesIO(buffer))
            content = self.to_md(doc, image_list, get_image_id_func())
            if len(image_list) > 0:
                content = content.replace('/api/image/', '/api/file/')
                save_image(image_list)
            return content
        except BaseException as e:
            traceback.print_exception(e)
            return f'{e}'
